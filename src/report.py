# -*- coding: utf-8 -*-
"""生成"链接及说明.docx"报告。

内容三段：① 分享链接(exsc URL) ② 说明文字 ③ 错点简述(B 给的一句话)。
仅在 B 判 A 错误时调用(答对的题不生成报告)。
"""
import pathlib
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
import config


def _set_cn_font(run, name="微软雅黑", size=11):
    """设置中英文字体(中文需 eastAsia)。"""
    run.font.size = Pt(size)
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def make_report(share_link, error_point, out_path, note_line=None):
    """生成报告 docx。share_link 为空时第一段写'(未获取到分享链接)'。"""
    note_line = note_line or config.REPORT_NOTE_LINE
    doc = Document()
    for text in (share_link or "(未获取到分享链接)", note_line, error_point or "(无明显错误)"):
        p = doc.add_paragraph()
        _set_cn_font(p.add_run(text))
    pathlib.Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
